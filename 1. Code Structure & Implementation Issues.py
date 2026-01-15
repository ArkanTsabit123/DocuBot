#1. Code Structure & Implementation Issues.py
"""
1. Code Structure & Implementation Issues
Missing utility functions (P1.3.4, P3.9.2, P1.15.2, P3.9.3)
Partial class implementations (P1.3.1, P1.3.2, P1.3.3, P1.5.1, P1.5.2)
Missing/extractor implementations (P1.4.1, P2.1.1, P2.1.3, P2.1.4)
Incomplete module methods (P1.6.1, P1.6.2, P1.6.3, P1.8.1, P1.9.1)
Missing configuration files (P1.2.2 can be enhanced)
"""

"""
DocuBot Code Structure & Implementation Completer
Fixes missing and incomplete implementations across the codebase
"""

import os
from pathlib import Path
import sys


class CodeCompleter:
    def __init__(self, project_dir="DocuBot"):
        self.project_dir = Path(project_dir).absolute()
        
    def complete_all_code(self):
        """Execute all code completion tasks"""
        print("=" * 60)
        print("DocuBot Code Completer")
        print("=" * 60)
        
        completions = [
            self.complete_text_cleaning,
            self.complete_document_processor,
            self.complete_pdf_extractor,
            self.complete_txt_extractor,
            self.complete_docx_extractor,
            self.complete_sqlite_client,
            self.complete_database_models,
            self.complete_chromadb_client,
            self.complete_llm_client,
            self.complete_embedding_service,
            self.complete_web_components,
            self.complete_cli_formatters,
            self.complete_app_config,
        ]
        
        for i, complete_func in enumerate(completions, 1):
            print(f"\n[{i}/{len(completions)}] Completing {complete_func.__name__}...")
            try:
                complete_func()
                print("   Completed")
            except Exception as e:
                print(f"   Error: {e}")
        
        print("\n" + "=" * 60)
        print("Code completion finished")
        print("=" * 60)
    
    def complete_text_cleaning(self):
        """Complete text cleaning utilities"""
        file_path = self.project_dir / "src" / "document_processing" / "cleaning.py"
        
        content = '''"""
Text Cleaning and Normalization Utilities
"""

import re
import unicodedata
from typing import List, Optional
import html


def clean_text(
    text: str,
    remove_html: bool = True,
    normalize_whitespace: bool = True,
    remove_special_chars: bool = False,
    to_lowercase: bool = False
) -> str:
    """
    Clean and normalize text.
    
    Args:
        text: Input text to clean
        remove_html: Remove HTML tags
        normalize_whitespace: Normalize whitespace characters
        remove_special_chars: Remove special characters
        to_lowercase: Convert to lowercase
        
    Returns:
        Cleaned text
    """
    if not text:
        return ""
    
    cleaned = text
    
    if remove_html:
        cleaned = html.unescape(cleaned)
        cleaned = re.sub(r'<[^>]+>', ' ', cleaned)
    
    cleaned = unicodedata.normalize('NFKC', cleaned)
    
    if normalize_whitespace:
        cleaned = re.sub(r'[\t\n\r\f\v]+', ' ', cleaned)
        cleaned = re.sub(r'\s+', ' ', cleaned)
        cleaned = cleaned.strip()
    
    if remove_special_chars:
        cleaned = re.sub(r'[^\w\s.,!?;:\-\'"()\[\]]', ' ', cleaned)
    
    if to_lowercase:
        cleaned = cleaned.lower()
    
    return cleaned


def split_into_sentences(text: str) -> List[str]:
    """
    Split text into sentences.
    
    Args:
        text: Input text
        
    Returns:
        List of sentences
    """
    sentences = re.split(r'[.!?]+\s*', text)
    sentences = [s.strip() for s in sentences if s.strip()]
    return sentences


def remove_stopwords(text: str, custom_stopwords: Optional[List[str]] = None) -> str:
    """
    Remove common stopwords from text.
    
    Args:
        text: Input text
        custom_stopwords: Custom list of stopwords
        
    Returns:
        Text with stopwords removed
    """
    default_stopwords = {
        'a', 'an', 'the', 'and', 'or', 'but', 'in', 'on', 'at', 'to',
        'for', 'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be',
        'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did',
        'will', 'would', 'shall', 'should', 'may', 'might', 'must',
        'can', 'could', 'i', 'you', 'he', 'she', 'it', 'we', 'they'
    }
    
    stopwords = default_stopwords
    if custom_stopwords:
        stopwords.update(custom_stopwords)
    
    words = text.split()
    filtered_words = [word for word in words if word.lower() not in stopwords]
    
    return ' '.join(filtered_words)


def normalize_numbers(text: str, replacement: str = "[NUM]") -> str:
    """
    Normalize numbers in text.
    
    Args:
        text: Input text
        replacement: String to replace numbers with
        
    Returns:
        Text with normalized numbers
    """
    normalized = re.sub(r'\b\d+\b', replacement, text)
    normalized = re.sub(r'\b\d+\.\d+\b', replacement, normalized)
    return normalized


def remove_extra_punctuation(text: str) -> str:
    """
    Remove excessive punctuation.
    
    Args:
        text: Input text
        
    Returns:
        Text with normalized punctuation
    """
    normalized = re.sub(r'([.!?])\1+', r'\1', text)
    normalized = re.sub(r'\s+([.,!?;:])', r'\1', normalized)
    normalized = re.sub(r'([.,!?;:])(?!\s|$)', r'\1 ', normalized)
    return normalized


def truncate_text(text: str, max_length: int = 1000) -> str:
    """
    Truncate text to maximum length while preserving word boundaries.
    
    Args:
        text: Input text
        max_length: Maximum character length
        
    Returns:
        Truncated text
    """
    if len(text) <= max_length:
        return text
    
    truncated = text[:max_length]
    last_space = truncated.rfind(' ')
    
    if last_space > max_length * 0.8:
        truncated = truncated[:last_space]
    
    return truncated + "..."


def calculate_readability_score(text: str) -> float:
    """
    Calculate simple readability score.
    
    Args:
        text: Input text
        
    Returns:
        Readability score (higher = easier to read)
    """
    sentences = split_into_sentences(text)
    words = text.split()
    
    if not sentences or not words:
        return 0.0
    
    words_per_sentence = len(words) / len(sentences)
    syllables_per_word = sum(len(re.findall(r'[aeiouy]+', word.lower())) 
                           for word in words) / len(words)
    
    score = 206.835 - (1.015 * words_per_sentence) - (84.6 * syllables_per_word)
    return max(0.0, min(100.0, score))


def clean_text_pipeline(text: str) -> str:
    """
    Complete text cleaning pipeline.
    
    Args:
        text: Input text
        
    Returns:
        Fully cleaned text
    """
    cleaned = clean_text(text, remove_html=True, normalize_whitespace=True)
    cleaned = remove_extra_punctuation(cleaned)
    cleaned = normalize_numbers(cleaned)
    return cleaned


if __name__ == "__main__":
    test_text = "Hello   World!! This is a test. 123 numbers. <b>HTML</b> tags."
    
    print("Original:", test_text)
    print("Cleaned:", clean_text_pipeline(test_text))
    print("Sentences:", split_into_sentences(test_text))
    print("Readability:", calculate_readability_score(test_text))
'''
        
        file_path.write_text(content)
    
    def complete_document_processor(self):
        """Complete document processor"""
        file_path = self.project_dir / "src" / "document_processing" / "processor.py"
        
        content = '''"""
Main Document Processing Pipeline
"""

import os
from pathlib import Path
from typing import Dict, List, Any, Optional
from datetime import datetime
import logging

from .extractors import create_extractor, get_supported_extensions
from .cleaning import clean_text_pipeline
from .chunking import chunk_text
from ..ai_engine.embedding_service import EmbeddingService
from ..vector_store.chroma_client import ChromaClient
from ..database.sqlite_client import SQLiteClient


logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Main document processing orchestrator"""
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        Initialize document processor.
        
        Args:
            config: Configuration dictionary
        """
        self.config = config or {
            'chunk_size': 500,
            'chunk_overlap': 50,
            'max_file_size_mb': 100,
            'supported_formats': ['.pdf', '.docx', '.txt', '.epub', '.md']
        }
        
        self.embedding_service = EmbeddingService()
        self.vector_store = ChromaClient()
        self.database = SQLiteClient()
        
        logger.info("DocumentProcessor initialized")
    
    def process_document(self, file_path: Path) -> Dict[str, Any]:
        """
        Complete document processing pipeline.
        
        Args:
            file_path: Path to document file
            
        Returns:
            Dictionary with processing results
        """
        start_time = datetime.now()
        results = {
            'success': False,
            'document_id': str(file_path),
            'error': None,
            'processing_time': 0,
            'chunks_created': 0
        }
        
        try:
            self._validate_file(file_path)
            logger.info(f"Processing document: {file_path.name}")
            
            extractor = create_extractor(file_path)
            if not extractor:
                raise ValueError(f"No extractor found for file: {file_path}")
            
            extraction_result = extractor.extract_with_metadata(file_path)
            if not extraction_result['success']:
                raise ValueError(f"Extraction failed: {extraction_result['error']}")
            
            raw_text = extraction_result['text']
            metadata = extraction_result['metadata']
            
            logger.info(f"Extracted {len(raw_text)} characters from {file_path.name}")
            
            cleaned_text = clean_text_pipeline(raw_text)
            logger.debug(f"Text cleaned: {len(cleaned_text)} characters remaining")
            
            chunks = chunk_text(
                text=cleaned_text,
                chunk_size=self.config['chunk_size'],
                chunk_overlap=self.config['chunk_overlap']
            )
            
            logger.info(f"Created {len(chunks)} chunks from document")
            
            chunk_texts = [chunk['text'] for chunk in chunks]
            embeddings = self.embedding_service.generate_embeddings(chunk_texts)
            
            vector_ids = self.vector_store.add_documents(
                texts=chunk_texts,
                embeddings=embeddings,
                metadatas=[{
                    **metadata,
                    'chunk_index': i,
                    'total_chunks': len(chunks),
                    'original_file': str(file_path),
                    'processed_at': datetime.now().isoformat()
                } for i in range(len(chunks))]
            )
            
            doc_id = self.database.add_document(
                file_path=str(file_path),
                file_name=file_path.name,
                file_type=metadata['file_extension'],
                file_size=metadata['file_size'],
                chunk_count=len(chunks),
                vector_ids=vector_ids,
                metadata=metadata
            )
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            results.update({
                'success': True,
                'document_id': doc_id,
                'chunks_created': len(chunks),
                'processing_time': processing_time,
                'metadata': metadata,
                'vector_ids': vector_ids
            })
            
            logger.info(f"Successfully processed {file_path.name} in {processing_time:.2f}s")
            
        except Exception as e:
            error_msg = f"Failed to process {file_path}: {str(e)}"
            logger.error(error_msg)
            results['error'] = error_msg
        
        return results
    
    def process_batch(self, file_paths: List[Path]) -> List[Dict[str, Any]]:
        """
        Process multiple documents.
        
        Args:
            file_paths: List of file paths
            
        Returns:
            List of processing results
        """
        results = []
        
        for file_path in file_paths:
            try:
                result = self.process_document(file_path)
                results.append(result)
            except Exception as e:
                logger.error(f"Batch processing failed for {file_path}: {e}")
                results.append({
                    'success': False,
                    'document_id': str(file_path),
                    'error': str(e)
                })
        
        return results
    
    def _validate_file(self, file_path: Path) -> bool:
        """
        Validate file before processing.
        
        Args:
            file_path: Path to validate
            
        Returns:
            True if valid
            
        Raises:
            ValueError: If file is invalid
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if not file_path.is_file():
            raise ValueError(f"Not a file: {file_path}")
        
        file_size = file_path.stat().st_size
        max_size = self.config.get('max_file_size_mb', 100) * 1024 * 1024
        
        if file_size > max_size:
            raise ValueError(f"File too large: {file_size} bytes (max: {max_size} bytes)")
        
        extension = file_path.suffix.lower()
        supported = self.config.get('supported_formats', get_supported_extensions())
        
        if extension not in supported:
            raise ValueError(f"Unsupported file format: {extension}")
        
        return True
    
    def get_supported_formats(self) -> List[str]:
        """
        Get list of supported file formats.
        
        Returns:
            List of supported file extensions
        """
        return self.config.get('supported_formats', get_supported_extensions())
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """
        Get processing statistics.
        
        Returns:
            Dictionary with statistics
        """
        return {
            'supported_formats': self.get_supported_formats(),
            'chunk_size': self.config.get('chunk_size', 500),
            'chunk_overlap': self.config.get('chunk_overlap', 50),
            'max_file_size_mb': self.config.get('max_file_size_mb', 100)
        }


_processor_instance = None

def get_processor(config: Optional[Dict[str, Any]] = None) -> DocumentProcessor:
    """
    Get or create DocumentProcessor instance.
    
    Args:
        config: Optional configuration
        
    Returns:
        DocumentProcessor instance
    """
    global _processor_instance
    
    if _processor_instance is None:
        _processor_instance = DocumentProcessor(config)
    
    return _processor_instance


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1:
        test_file = Path(sys.argv[1])
        if test_file.exists():
            processor = DocumentProcessor()
            result = processor.process_document(test_file)
            print("Processing result:", result)
        else:
            print(f"File not found: {test_file}")
    else:
        print("Usage: python processor.py <file_path>")
        print("Supported formats:", get_supported_extensions())
'''
        
        file_path.write_text(content)
    
    def complete_pdf_extractor(self):
        """Complete PDF extractor"""
        file_path = self.project_dir / "src" / "document_processing" / "extractors" / "pdf_extractor.py"
        
        content = '''"""
PDF Document Extractor
"""

from pathlib import Path
from typing import Dict, Any
import logging

from .base_extractor import BaseExtractor, TextExtractor

try:
    import PyPDF2
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    logging.warning("PyPDF2 or pdfplumber not installed. PDF extraction disabled.")


logger = logging.getLogger(__name__)


class PDFExtractor(BaseExtractor):
    """PDF document extractor"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.pdf']
        self.name = 'PDFExtractor'
        
        if not PDF_SUPPORT:
            raise ImportError("Required packages not installed: PyPDF2, pdfplumber")
    
    def extract(self, file_path: Path) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        self.validate_file(file_path)
        
        extracted_text = ""
        
        try:
            extracted_text = self._extract_with_pdfplumber(file_path)
            
            if len(extracted_text.strip()) < 100:
                logger.debug(f"pdfplumber extracted only {len(extracted_text)} chars, trying PyPDF2")
                pyPDF_text = self._extract_with_pypdf2(file_path)
                
                if len(pyPDF_text) > len(extracted_text):
                    extracted_text = pyPDF_text
            
        except Exception as e:
            logger.error(f"PDF extraction failed with pdfplumber: {e}")
            extracted_text = self._extract_with_pypdf2(file_path)
        
        if not extracted_text.strip():
            raise ValueError(f"No text could be extracted from PDF: {file_path}")
        
        logger.info(f"Extracted {len(extracted_text)} characters from PDF: {file_path.name}")
        return extracted_text
    
    def _extract_with_pdfplumber(self, file_path: Path) -> str:
        """
        Extract text using pdfplumber.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        text_parts = []
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    page_text = page.extract_text()
                    
                    if page_text:
                        text_parts.append(f"--- Page {page_num} ---\n{page_text}\n")
                    
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            table_text = self._format_table(table)
                            text_parts.append(f"\nTable on page {page_num}:\n{table_text}\n")
                            
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num}: {e}")
                    continue
        
        return "\n\n".join(text_parts)
    
    def _extract_with_pypdf2(self, file_path: Path) -> str:
        """
        Extract text using PyPDF2.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        text_parts = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num in range(len(pdf_reader.pages)):
                try:
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    
                    if page_text:
                        text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}\n")
                        
                except Exception as e:
                    logger.warning(f"PyPDF2 error on page {page_num + 1}: {e}")
                    continue
        
        return "\n\n".join(text_parts)
    
    def _format_table(self, table_data: list) -> str:
        """
        Format table data as readable text.
        
        Args:
            table_data: 2D list of table cells
            
        Returns:
            Formatted table text
        """
        if not table_data:
            return ""
        
        formatted_lines = []
        
        for row in table_data:
            row_cells = [str(cell) if cell is not None else "" for cell in row]
            formatted_lines.append(" | ".join(row_cells))
        
        return "\n".join(formatted_lines)
    
    def get_metadata(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract PDF-specific metadata.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Dictionary with metadata
        """
        metadata = super().get_metadata(file_path)
        
        if PDF_SUPPORT:
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    
                    pdf_info = pdf_reader.metadata
                    if pdf_info:
                        metadata.update({
                            'pdf_title': pdf_info.get('/Title', ''),
                            'pdf_author': pdf_info.get('/Author', ''),
                            'pdf_subject': pdf_info.get('/Subject', ''),
                            'pdf_keywords': pdf_info.get('/Keywords', ''),
                            'pdf_creator': pdf_info.get('/Creator', ''),
                            'pdf_producer': pdf_info.get('/Producer', ''),
                            'pdf_creation_date': pdf_info.get('/CreationDate', ''),
                            'pdf_modification_date': pdf_info.get('/ModDate', ''),
                        })
                    
                    metadata['page_count'] = len(pdf_reader.pages)
                    
            except Exception as e:
                logger.warning(f"Could not extract PDF metadata: {e}")
        
        return metadata


if PDF_SUPPORT:
    from ..extractors import register_extractor
    register_extractor(PDFExtractor, ['.pdf'])


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1:
        extractor = PDFExtractor()
        result = extractor.extract_with_metadata(Path(sys.argv[1]))
        print("Extraction result:", {
            'success': result['success'],
            'text_length': len(result['text']),
            'metadata': result['metadata']
        })
    else:
        print("Usage: python pdf_extractor.py <pdf_file>")
'''
        
        file_path.write_text(content)
    
    def complete_txt_extractor(self):
        """Complete TXT extractor"""
        file_path = self.project_dir / "src" / "document_processing" / "extractors" / "txt_extractor.py"
        
        content = '''"""
Plain Text Document Extractor
"""

from pathlib import Path
from typing import Dict, Any
import logging

from .base_extractor import TextExtractor

logger = logging.getLogger(__name__)


class TXTExtractor(TextExtractor):
    """Plain text document extractor"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.txt', '.text', '.log', '.csv', '.tsv']
        self.name = 'TXTExtractor'
    
    def extract(self, file_path: Path) -> str:
        """
        Extract text from plain text file.
        
        Args:
            file_path: Path to text file
            
        Returns:
            Extracted text
        """
        self.validate_file(file_path)
        
        encoding = self.detect_encoding(file_path)
        
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                text = f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        
        logger.info(f"Extracted {len(text)} characters from text file: {file_path.name}")
        return text
    
    def get_metadata(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract text file metadata.
        
        Args:
            file_path: Path to text file
            
        Returns:
            Dictionary with metadata
        """
        metadata = super().get_metadata(file_path)
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
                lines = content.split('\n')
                words = content.split()
                
                metadata.update({
                    'line_count': len(lines),
                    'word_count': len(words),
                    'character_count': len(content),
                    'encoding': self.detect_encoding(file_path)
                })
                
                preview_lines = lines[:5]
                metadata['preview'] = '\n'.join(preview_lines)
                
        except Exception as e:
            logger.warning(f"Could not analyze text file: {e}")
        
        return metadata


from ..extractors import register_extractor
register_extractor(TXTExtractor, ['.txt', '.text', '.log'])


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1:
        extractor = TXTExtractor()
        result = extractor.extract_with_metadata(Path(sys.argv[1]))
        print("Extraction result:", {
            'success': result['success'],
            'text_length': len(result['text']),
            'metadata': result['metadata']
        })
    else:
        print("Usage: python txt_extractor.py <text_file>")
'''
        
        file_path.write_text(content)
    
    def complete_docx_extractor(self):
        """Complete DOCX extractor"""
        file_path = self.project_dir / "src" / "document_processing" / "extractors" / "docx_extractor.py"
        
        content = '''"""
Microsoft Word DOCX Document Extractor
"""

from pathlib import Path
from typing import Dict, Any
import logging

from .base_extractor import BaseExtractor

try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    logging.warning("python-docx not installed. DOCX extraction disabled.")


logger = logging.getLogger(__name__)


class DOCXExtractor(BaseExtractor):
    """Microsoft Word DOCX document extractor"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.docx', '.doc']
        self.name = 'DOCXExtractor'
        
        if not DOCX_SUPPORT:
            raise ImportError("Required package not installed: python-docx")
    
    def extract(self, file_path: Path) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        self.validate_file(file_path)
        
        try:
            doc = docx.Document(file_path)
            
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    text_parts.append(f"\n{table_text}\n")
            
            extracted_text = "\n\n".join(text_parts)
            
            if not extracted_text.strip():
                raise ValueError(f"No text found in DOCX file: {file_path}")
            
            logger.info(f"Extracted {len(extracted_text)} characters from DOCX: {file_path.name}")
            return extracted_text
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise
    
    def _extract_table_text(self, table) -> str:
        """
        Extract text from a table.
        
        Args:
            table: docx table object
            
        Returns:
            Formatted table text
        """
        table_text = []
        
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_cells.append(cell_text)
            
            if row_cells:
                table_text.append(" | ".join(row_cells))
        
        return "\n".join(table_text)
    
    def get_metadata(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract DOCX-specific metadata.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary with metadata
        """
        metadata = super().get_metadata(file_path)
        
        if DOCX_SUPPORT:
            try:
                doc = docx.Document(file_path)
                core_props = doc.core_properties
                
                metadata.update({
                    'docx_title': core_props.title or '',
                    'docx_author': core_props.author or '',
                    'docx_subject': core_props.subject or '',
                    'docx_keywords': core_props.keywords or '',
                    'docx_comments': core_props.comments or '',
                    'docx_created': core_props.created.isoformat() if core_props.created else '',
                    'docx_modified': core_props.modified.isoformat() if core_props.modified else '',
                    'docx_category': core_props.category or '',
                    'docx_paragraph_count': len(doc.paragraphs),
                    'docx_table_count': len(doc.tables),
                })
                
                word_count = sum(len(para.text.split()) for para in doc.paragraphs)
                metadata['docx_word_count'] = word_count
                
                preview = ""
                for para in doc.paragraphs[:3]:
                    if para.text.strip():
                        preview += para.text + "\n"
                metadata['preview'] = preview.strip()
                
            except Exception as e:
                logger.warning(f"Could not extract DOCX metadata: {e}")
        
        return metadata


if DOCX_SUPPORT:
    from ..extractors import register_extractor
    register_extractor(DOCXExtractor, ['.docx', '.doc'])


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1:
        extractor = DOCXExtractor()
        result = extractor.extract_with_metadata(Path(sys.argv[1]))
        print("Extraction result:", {
            'success': result['success'],
            'text_length': len(result['text']),
            'metadata': result['metadata']
        })
    else:
        print("Usage: python docx_extractor.py <docx_file>")
'''
        
        file_path.write_text(content)
    
    def complete_sqlite_client(self):
        """Complete SQLite client"""
        file_path = self.project_dir / "src" / "database" / "sqlite_client.py"
        
        content = '''"""
SQLite Database Client with CRUD Operations
"""

import sqlite3
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime
import logging
import json

logger = logging.getLogger(__name__)


class SQLiteClient:
    """SQLite database client"""
    
    def __init__(self, db_path: Optional[Path] = None):
        """
        Initialize SQLite client.
        
        Args:
            db_path: Path to SQLite database file
        """
        if db_path is None:
            from ...core.constants import DATABASE_DIR, DATABASE_NAME
            db_path = DATABASE_DIR / DATABASE_NAME
        
        self.db_path = db_path
        self.db_path.parent.mkdir(parents=True, exist_ok=True)
        
        self._init_database()
        
        logger.info(f"SQLite client initialized: {self.db_path}")
    
    def _init_database(self):
        """Initialize database schema"""
        conn = self._get_connection()
        cursor = conn.cursor()
        
        tables_sql = self._get_table_definitions()
        
        for table_name, create_sql in tables_sql.items():
            try:
                cursor.execute(f"SELECT name FROM sqlite_master WHERE type='table' AND name='{table_name}'")
                if not cursor.fetchone():
                    cursor.execute(create_sql)
                    logger.debug(f"Created table: {table_name}")
            except Exception as e:
                logger.error(f"Error creating table {table_name}: {e}")
        
        conn.commit()
        conn.close()
    
    def _get_table_definitions(self) -> Dict[str, str]:
        """Get SQL table definitions"""
        return {
            'documents': """
                CREATE TABLE documents (
                    id TEXT PRIMARY KEY,
                    file_path TEXT NOT NULL,
                    file_name TEXT NOT NULL,
                    file_type TEXT NOT NULL,
                    file_size INTEGER,
                    upload_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    processing_status TEXT DEFAULT 'pending',
                    processing_error TEXT,
                    metadata_json TEXT,
                    vector_ids_json TEXT,
                    chunk_count INTEGER DEFAULT 0,
                    word_count INTEGER DEFAULT 0,
                    language TEXT,
                    tags_json TEXT,
                    summary TEXT,
                    is_indexed BOOLEAN DEFAULT FALSE,
                    indexed_at TIMESTAMP,
                    last_accessed TIMESTAMP,
                    access_count INTEGER DEFAULT 0
                )
            """,
            'chunks': """
                CREATE TABLE chunks (
                    id TEXT PRIMARY KEY,
                    document_id TEXT NOT NULL,
                    chunk_index INTEGER NOT NULL,
                    text_content TEXT NOT NULL,
                    cleaned_text TEXT NOT NULL,
                    token_count INTEGER,
                    embedding_model TEXT,
                    vector_id TEXT NOT NULL,
                    metadata_json TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (document_id) REFERENCES documents(id) ON DELETE CASCADE
                )
            """,
            'conversations': """
                CREATE TABLE conversations (
                    id TEXT PRIMARY KEY,
                    title TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    message_count INTEGER DEFAULT 0,
                    total_tokens INTEGER DEFAULT 0,
                    tags_json TEXT,
                    is_archived BOOLEAN DEFAULT FALSE,
                    export_path TEXT
                )
            """,
            'messages': """
                CREATE TABLE messages (
                    id TEXT PRIMARY KEY,
                    conversation_id TEXT NOT NULL,
                    role TEXT NOT NULL,
                    content TEXT NOT NULL,
                    tokens INTEGER,
                    model_used TEXT,
                    sources_json TEXT,
                    processing_time_ms INTEGER,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (conversation_id) REFERENCES conversations(id) ON DELETE CASCADE
                )
            """,
            'settings': """
                CREATE TABLE settings (
                    key TEXT PRIMARY KEY,
                    value TEXT,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """
        }
    
    def _get_connection(self) -> sqlite3.Connection:
        """Get database connection"""
        conn = sqlite3.connect(self.db_path)
        conn.row_factory = sqlite3.Row
        return conn
    
    def add_document(self, 
                    file_path: str,
                    file_name: str,
                    file_type: str,
                    file_size: int,
                    chunk_count: int = 0,
                    vector_ids: Optional[List[str]] = None,
                    metadata: Optional[Dict[str, Any]] = None) -> str:
        """
        Add a new document.
        
        Args:
            file_path: Full path to document
            file_name: Document file name
            file_type: File extension/type
            file_size: File size in bytes
            chunk_count: Number of chunks
            vector_ids: List of vector IDs
            metadata: Additional metadata
            
        Returns:
            Document ID
        """
        import uuid
        
        doc_id = str(uuid.uuid4())
        now = datetime.now().isoformat()
        
        conn = self._get_connection()
        cursor = conn.cursor()
        
        cursor.execute("""
            INSERT INTO documents 
            (id, file_path, file_name, file_type, file_size, chunk_count, 
             vector_ids_json, metadata_json, last_accessed)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            doc_id,
            file_path,
            file_name,
            file_type,
            file_size,
            chunk_count,
            json.dumps(vector_ids or []),
            json.dumps(metadata or {}),
            now
        ))
        
        conn.commit()
        conn.close()
        
        logger.info(f"Added document: {file_name} (ID: {doc_id})")
        return doc_id
    
    def get_document(self, doc_id: str) -> Optional[Dict[str, Any]]:
        """
        Get document by ID.
        
        Args:
            doc_id: Document ID
            
        Returns:
            Document data or None
        """
        conn = self._get_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT * FROM documents WHERE id = ?", (doc_id,))
        row = cursor.fetchone()
        
        conn.close()
        
        if row:
            return self._row_to_dict(row)
        return None
    
    def update_document(self, doc_id: str, updates: Dict[str, Any]) -> bool:
        """
        Update document.
        
        Args:
            doc_id: Document ID
            updates: Fields to update
            
        Returns:
            True if successful
        """
        if not updates:
            return False
        
        set_clause = []
        values = []
        
        for key, value in updates.items():
            if key in ['vector_ids_json', 'metadata_json', 'tags_json']:
                value = json.dumps(value) if value else '[]'
            
            set_clause.append(f"{key} = ?")
            values.append(value)
        
        values.append(doc_id)
        
        sql = f"UPDATE documents SET {', '.join(set_clause)} WHERE id = ?"
        
        conn = self._get_connection()
        cursor = conn.cursor()
        
        try:
            cursor.execute(sql, values)
            conn.commit()
            success = cursor.rowcount > 0
        except Exception as e:
            logger.error(f"Error updating document {doc_id}: {e}")
            success = False
        finally:
            conn.close()
        
        return success
    
    def delete_document(self, doc_id: str) -> bool:
        """
        Delete document.
        
        Args:
            doc_id: Document ID
            
        Returns:
            True if successful
        """
        conn = self._get_connection()
        cursor = conn.cursor()
        
        try:
            cursor.execute("DELETE FROM documents WHERE id = ?", (doc_id,))
            conn.commit()
            success = cursor.rowcount > 0
            
            if success:
                logger.info(f"Deleted document: {doc_id}")
        except Exception as e:
            logger.error(f"Error deleting document {doc_id}: {e}")
            success = False
        finally:
            conn.close()
        
        return success
    
    def list_documents(self, 
                      limit: int = 100,
                      offset: int = 0,
                      status: Optional[str] = None) -> List[Dict[str, Any]]:
        """
        List documents.
        
        Args:
            limit: Maximum number of documents
            offset: Offset for pagination
            status: Filter by processing status
            
        Returns:
            List of documents
        """
        conn = self._get_connection()
        cursor = conn.cursor()
        
        where_clause = "WHERE 1=1"
        params = []
        
        if status:
            where_clause += " AND processing_status = ?"
            params.append(status)
        
        sql = f"""
            SELECT * FROM documents 
            {where_clause}
            ORDER BY upload_date DESC
            LIMIT ? OFFSET ?
        """
        
        params.extend([limit, offset])
        
        cursor.execute(sql, params)
        rows = cursor.fetchall()
        conn.close()
        
        return [self._row_to_dict(row) for row in rows]
    
    def add_chunk(self,
                 document_id: str,
                 chunk_index: int,
                 text_content: str,
                 cleaned_text: str,
                 vector_id: str,
                 token_count: Optional[int] = None,
                 embedding_model: Optional[str] = None,
                 metadata: Optional[Dict[str, Any]] = None) -> str:
        """
        Add a chunk.
        
        Args:
            document_id: Parent document ID
            chunk_index: Chunk index in document
            text_content: Original text content
            cleaned_text: Cleaned text content
            vector_id: Vector store ID
            token_count: Number of tokens
            embedding_model: Embedding model used
            metadata: Additional metadata
            
        Returns:
            Chunk ID
        """
        import uuid
        
        chunk_id = str(uuid.uuid4())
        
        conn = self._get_connection()
        cursor = conn.cursor()
        
        cursor.execute("""
            INSERT INTO chunks 
            (id, document_id, chunk_index, text_content, cleaned_text, 
             token_count, embedding_model, vector_id, metadata_json)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            chunk_id,
            document_id,
            chunk_index,
            text_content,
            cleaned_text,
            token_count,
            embedding_model,
            vector_id,
            json.dumps(metadata or {})
        ))
        
        conn.commit()
        conn.close()
        
        return chunk_id
    
    def get_chunks_by_document(self, document_id: str) -> List[Dict[str, Any]]:
        """
        Get all chunks for a document.
        
        Args:
            document_id: Document ID
            
        Returns:
            List of chunks
        """
        conn = self._get_connection()
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT * FROM chunks 
            WHERE document_id = ?
            ORDER BY chunk_index
        """, (document_id,))
        
        rows = cursor.fetchall()
        conn.close()
        
        return [self._row_to_dict(row) for row in rows]
    
    def _row_to_dict(self, row: sqlite3.Row) -> Dict[str, Any]:
        """Convert SQLite row to dictionary"""
        result = dict(row)
        
        json_fields = ['vector_ids_json', 'metadata_json', 'tags_json', 'sources_json']
        
        for field in json_fields:
            if field in result and result[field]:
                try:
                    result[field.replace('_json', '')] = json.loads(result[field])
                except (json.JSONDecodeError, TypeError):
                    result[field.replace('_json', '')] = []
                finally:
                    del result[field]
        
        return result
    
    def execute_query(self, sql: str, params: Tuple = ()) -> List[Dict[str, Any]]:
        """
        Execute custom SQL query.
        
        Args:
            sql: SQL query
            params: Query parameters
            
        Returns:
            Query results
        """
        conn = self._get_connection()
        cursor = conn.cursor()
        
        try:
            cursor.execute(sql, params)
            rows = cursor.fetchall()
            results = [self._row_to_dict(row) for row in rows]
        except Exception as e:
            logger.error(f"Query execution error: {e}")
            results = []
        finally:
            conn.close()
        
        return results
    
    def get_stats(self) -> Dict[str, Any]:
        """
        Get database statistics.
        
        Returns:
            Statistics dictionary
        """
        stats = {}
        
        try:
            conn = self._get_connection()
            cursor = conn.cursor()
            
            cursor.execute("""
                SELECT processing_status, COUNT(*) as count
                FROM documents
                GROUP BY processing_status
            """)
            stats['documents_by_status'] = dict(cursor.fetchall())
            
            cursor.execute("SELECT COUNT(*) FROM documents")
            stats['total_documents'] = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM chunks")
            stats['total_chunks'] = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM conversations")
            stats['total_conversations'] = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM messages")
            stats['total_messages'] = cursor.fetchone()[0]
            
            db_size = self.db_path.stat().st_size if self.db_path.exists() else 0
            stats['database_size_bytes'] = db_size
            
            conn.close()
            
        except Exception as e:
            logger.error(f"Error getting stats: {e}")
            stats['error'] = str(e)
        
        return stats


_db_instance = None

def get_database_client(db_path: Optional[Path] = None) -> SQLiteClient:
    """
    Get or create SQLiteClient instance.
    
    Args:
        db_path: Optional database path
        
    Returns:
        SQLiteClient instance
    """
    global _db_instance
    
    if _db_instance is None:
        _db_instance = SQLiteClient(db_path)
    
    return _db_instance


if __name__ == "__main__":
    client = SQLiteClient()
    
    doc_id = client.add_document(
        file_path="/test/document.pdf",
        file_name="test.pdf",
        file_type=".pdf",
        file_size=1024,
        chunk_count=5,
        metadata={"test": "data"}
    )
    
    print(f"Added document with ID: {doc_id}")
    
    doc = client.get_document(doc_id)
    print(f"Retrieved document: {doc['file_name']}")
    
    docs = client.list_documents(limit=5)
    print(f"Total documents: {len(docs)}")
    
    stats = client.get_stats()
    print(f"Database stats: {stats}")
'''
        
        file_path.write_text(content)
    
    def complete_database_models(self):
        """Complete database models"""
        file_path = self.project_dir / "src" / "database" / "models.py"
        
        content = '''"""
SQLAlchemy ORM Models for DocuBot Database
"""

from datetime import datetime
from typing import List, Optional, Dict, Any
from sqlalchemy import create_engine, Column, Integer, String, Text, Boolean, DateTime, ForeignKey
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import relationship, sessionmaker, Session
from sqlalchemy import JSON
import uuid

from ..core.constants import DATABASE_DIR, DATABASE_NAME

Base = declarative_base()


class Document(Base):
    """Document model representing uploaded files"""
    
    __tablename__ = 'documents'
    
    id = Column(String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    file_path = Column(Text, nullable=False)
    file_name = Column(String(255), nullable=False)
    file_type = Column(String(10), nullable=False)
    file_size = Column(Integer)
    
    upload_date = Column(DateTime, default=datetime.utcnow)
    processing_status = Column(String(20), default='pending')
    processing_error = Column(Text)
    
    metadata_json = Column(JSON, default=dict)
    vector_ids_json = Column(JSON, default=list)
    
    chunk_count = Column(Integer, default=0)
    word_count = Column(Integer, default=0)
    language = Column(String(10))
    
    tags_json = Column(JSON, default=list)
    summary = Column(Text)
    
    is_indexed = Column(Boolean, default=False)
    indexed_at = Column(DateTime)
    
    last_accessed = Column(DateTime, default=datetime.utcnow)
    access_count = Column(Integer, default=0)
    
    chunks = relationship("Chunk", back_populates="document", cascade="all, delete-orphan")
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert model to dictionary"""
        return {
            'id': self.id,
            'file_path': self.file_path,
            'file_name': self.file_name,
            'file_type': self.file_type,
            'file_size': self.file_size,
            'upload_date': self.upload_date.isoformat() if self.upload_date else None,
            'processing_status': self.processing_status,
            'processing_error': self.processing_error,
            'metadata': self.metadata_json,
            'vector_ids': self.vector_ids_json,
            'chunk_count': self.chunk_count,
            'word_count': self.word_count,
            'language': self.language,
            'tags': self.tags_json,
            'summary': self.summary,
            'is_indexed': self.is_indexed,
            'indexed_at': self.indexed_at.isoformat() if self.indexed_at else None,
            'last_accessed': self.last_accessed.isoformat() if self.last_accessed else None,
            'access_count': self.access_count
        }
    
    def update_last_accessed(self):
        """Update last accessed timestamp"""
        self.last_accessed = datetime.utcnow()
        self.access_count += 1


class Chunk(Base):
    """Text chunk model"""
    
    __tablename__ = 'chunks'
    
    id = Column(String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    document_id = Column(String(36), ForeignKey('documents.id', ondelete='CASCADE'), nullable=False)
    chunk_index = Column(Integer, nullable=False)
    
    text_content = Column(Text, nullable=False)
    cleaned_text = Column(Text, nullable=False)
    token_count = Column(Integer)
    
    embedding_model = Column(String(50))
    vector_id = Column(String(255), nullable=False)
    
    metadata_json = Column(JSON, default=dict)
    created_at = Column(DateTime, default=datetime.utcnow)
    
    document = relationship("Document", back_populates="chunks")
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert model to dictionary"""
        return {
            'id': self.id,
            'document_id': self.document_id,
            'chunk_index': self.chunk_index,
            'text_content': self.text_content,
            'cleaned_text': self.cleaned_text,
            'token_count': self.token_count,
            'embedding_model': self.embedding_model,
            'vector_id': self.vector_id,
            'metadata': self.metadata_json,
            'created_at': self.created_at.isoformat() if self.created_at else None
        }


class Conversation(Base):
    """Conversation model"""
    
    __tablename__ = 'conversations'
    
    id = Column(String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    title = Column(String(255))
    
    created_at = Column(DateTime, default=datetime.utcnow)
    updated_at = Column(DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    message_count = Column(Integer, default=0)
    total_tokens = Column(Integer, default=0)
    
    tags_json = Column(JSON, default=list)
    is_archived = Column(Boolean, default=False)
    export_path = Column(Text)
    
    messages = relationship("Message", back_populates="conversation", cascade="all, delete-orphan")
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert model to dictionary"""
        return {
            'id': self.id,
            'title': self.title,
            'created_at': self.created_at.isoformat() if self.created_at else None,
            'updated_at': self.updated_at.isoformat() if self.updated_at else None,
            'message_count': self.message_count,
            'total_tokens': self.total_tokens,
            'tags': self.tags_json,
            'is_archived': self.is_archived,
            'export_path': self.export_path
        }
    
    def update_message_count(self):
        """Update message count"""
        self.message_count = len(self.messages)
        self.total_tokens = sum(msg.tokens or 0 for msg in self.messages)


class Message(Base):
    """Chat message model"""
    
    __tablename__ = 'messages'
    
    id = Column(String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    conversation_id = Column(String(36), ForeignKey('conversations.id', ondelete='CASCADE'), nullable=False)
    
    role = Column(String(20), nullable=False)
    content = Column(Text, nullable=False)
    tokens = Column(Integer)
    
    model_used = Column(String(50))
    sources_json = Column(JSON, default=list)
    processing_time_ms = Column(Integer)
    
    created_at = Column(DateTime, default=datetime.utcnow)
    
    conversation = relationship("Conversation", back_populates="messages")
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert model to dictionary"""
        return {
            'id': self.id,
            'conversation_id': self.conversation_id,
            'role': self.role,
            'content': self.content,
            'tokens': self.tokens,
            'model_used': self.model_used,
            'sources': self.sources_json,
            'processing_time_ms': self.processing_time_ms,
            'created_at': self.created_at.isoformat() if self.created_at else None
        }


class Setting(Base):
    """Application settings model"""
    
    __tablename__ = 'settings'
    
    key = Column(String(100), primary_key=True)
    value = Column(Text)
    updated_at = Column(DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert model to dictionary"""
        return {
            'key': self.key,
            'value': self.value,
            'updated_at': self.updated_at.isoformat() if self.updated_at else None
        }


_engine = None
_SessionLocal = None

def get_engine():
    """Get SQLAlchemy engine"""
    global _engine
    
    if _engine is None:
        db_url = f"sqlite:///{DATABASE_DIR / DATABASE_NAME}"
        _engine = create_engine(db_url, connect_args={"check_same_thread": False})
    
    return _engine

def get_session_local():
    """Get session factory"""
    global _SessionLocal
    
    if _SessionLocal is None:
        _SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=get_engine())
    
    return _SessionLocal

def get_session() -> Session:
    """Get database session"""
    SessionLocal = get_session_local()
    return SessionLocal()

def create_tables():
    """Create all database tables"""
    engine = get_engine()
    Base.metadata.create_all(bind=engine)

def init_database():
    """Initialize database with tables"""
    DATABASE_DIR.mkdir(parents=True, exist_ok=True)
    create_tables()
    
    session = get_session()
    try:
        existing = session.query(Setting).count()
        
        if existing == 0:
            default_settings = [
                Setting(key="app.version", value="1.0.0"),
                Setting(key="app.default_chunk_size", value="500"),
                Setting(key="app.default_chunk_overlap", value="50"),
                Setting(key="app.default_llm_model", value="llama2:7b"),
                Setting(key="app.default_embedding_model", value="all-MiniLM-L6-v2"),
            ]
            
            session.add_all(default_settings)
            session.commit()
            print("Database initialized with default settings")
        else:
            print(f"Database already initialized with {existing} settings")
            
    except Exception as e:
        session.rollback()
        print(f"Error initializing database: {e}")
    finally:
        session.close()


if __name__ == "__main__":
    init_database()
    
    session = get_session()
    
    try:
        test_doc = Document(
            file_path="/test/example.pdf",
            file_name="example.pdf",
            file_type=".pdf",
            file_size=1024,
            processing_status="completed",
            chunk_count=5,
            word_count=1000
        )
        
        session.add(test_doc)
        session.commit()
        
        print(f"Created test document: {test_doc.id}")
        
        doc = session.query(Document).filter_by(file_name="example.pdf").first()
        print(f"Retrieved document: {doc.to_dict()}")
        
        chunk = Chunk(
            document_id=doc.id,
            chunk_index=0,
            text_content="This is a test chunk.",
            cleaned_text="This is a test chunk.",
            token_count=5,
            vector_id="test_vector_123"
        )
        
        session.add(chunk)
        session.commit()
        
        print(f"Created test chunk: {chunk.id}")
        
        doc_chunks = doc.chunks
        print(f"Document has {len(doc_chunks)} chunks")
        
    except Exception as e:
        print(f"Error: {e}")
        session.rollback()
    finally:
        session.close()
'''
        
        file_path.write_text(content)
    
    def complete_chromadb_client(self):
        """Complete ChromaDB client"""
        file_path = self.project_dir / "src" / "vector_store" / "chroma_client.py"
        
        content = '''"""
ChromaDB Vector Store Client for DocuBot
"""

import chromadb
from chromadb.config import Settings
from pathlib import Path
from typing import List, Dict, Any, Optional
import logging
from datetime import datetime
import uuid

from ..core.constants import DATABASE_DIR, VECTOR_DB_NAME

logger = logging.getLogger(__name__)


class ChromaClient:
    """ChromaDB vector store client"""
    
    def __init__(self, persist_directory: Optional[Path] = None, collection_name: str = "documents"):
        """
        Initialize ChromaDB client.
        
        Args:
            persist_directory: Directory to persist ChromaDB data
            collection_name: Name of the collection to use
        """
        if persist_directory is None:
            persist_directory = DATABASE_DIR / VECTOR_DB_NAME
        
        self.persist_directory = persist_directory
        self.collection_name = collection_name
        
        self.persist_directory.mkdir(parents=True, exist_ok=True)
        
        self.client = chromadb.PersistentClient(
            path=str(self.persist_directory),
            settings=Settings(
                anonymized_telemetry=False,
                allow_reset=True
            )
        )
        
        self.collection = self._get_or_create_collection()
        
        logger.info(f"ChromaDB client initialized: {self.persist_directory}")
    
    def _get_or_create_collection(self) -> chromadb.Collection:
        """
        Get existing collection or create new one.
        
        Returns:
            ChromaDB collection
        """
        try:
            collection = self.client.get_collection(name=self.collection_name)
            logger.debug(f"Using existing collection: {self.collection_name}")
            return collection
        except ValueError:
            collection = self.client.create_collection(
                name=self.collection_name,
                metadata={"description": "DocuBot document embeddings"}
            )
            logger.info(f"Created new collection: {self.collection_name}")
            return collection
    
    def add_documents(self,
                     texts: List[str],
                     embeddings: Optional[List[List[float]]] = None,
                     metadatas: Optional[List[Dict[str, Any]]] = None,
                     ids: Optional[List[str]] = None) -> List[str]:
        """
        Add documents to vector store.
        
        Args:
            texts: List of text documents
            embeddings: Optional pre-computed embeddings
            metadatas: Optional metadata for each document
            ids: Optional custom IDs for documents
            
        Returns:
            List of document IDs
        """
        if not texts:
            return []
        
        if ids is None:
            ids = [str(uuid.uuid4()) for _ in range(len(texts))]
        
        if metadatas is None:
            metadatas = [{} for _ in range(len(texts))]
        
        timestamp = datetime.now().isoformat()
        for metadata in metadatas:
            metadata['added_at'] = metadata.get('added_at', timestamp)
        
        try:
            if embeddings:
                self.collection.add(
                    embeddings=embeddings,
                    documents=texts,
                    metadatas=metadatas,
                    ids=ids
                )
            else:
                self.collection.add(
                    documents=texts,
                    metadatas=metadatas,
                    ids=ids
                )
            
            logger.info(f"Added {len(texts)} documents to vector store")
            return ids
            
        except Exception as e:
            logger.error(f"Error adding documents: {e}")
            raise
    
    def search(self,
              query: str,
              n_results: int = 5,
              where: Optional[Dict[str, Any]] = None,
              where_document: Optional[Dict[str, Any]] = None,
              include: List[str] = None) -> Dict[str, Any]:
        """
        Search for similar documents.
        
        Args:
            query: Search query text
            n_results: Number of results to return
            where: Filter by metadata
            where_document: Filter by document content
            include: What to include in results
            
        Returns:
            Search results
        """
        if include is None:
            include = ["documents", "metadatas", "distances"]
        
        try:
            results = self.collection.query(
                query_texts=[query],
                n_results=n_results,
                where=where,
                where_document=where_document,
                include=include
            )
            
            formatted_results = {
                'ids': results['ids'][0] if results['ids'] else [],
                'documents': results['documents'][0] if results['documents'] else [],
                'metadatas': results['metadatas'][0] if results['metadatas'] else [],
                'distances': results['distances'][0] if results['distances'] else [],
                'count': len(results['ids'][0]) if results['ids'] else 0
            }
            
            logger.debug(f"Search returned {formatted_results['count']} results")
            return formatted_results
            
        except Exception as e:
            logger.error(f"Error searching: {e}")
            return {
                'ids': [],
                'documents': [],
                'metadatas': [],
                'distances': [],
                'count': 0,
                'error': str(e)
            }
    
    def search_with_embeddings(self,
                              query_embeddings: List[List[float]],
                              n_results: int = 5,
                              where: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """
        Search using pre-computed embeddings.
        
        Args:
            query_embeddings: Query embeddings
            n_results: Number of results to return
            where: Filter by metadata
            
        Returns:
            Search results
        """
        try:
            results = self.collection.query(
                query_embeddings=query_embeddings,
                n_results=n_results,
                where=where,
                include=["documents", "metadatas", "distances"]
            )
            
            formatted_results = {
                'ids': results['ids'][0] if results['ids'] else [],
                'documents': results['documents'][0] if results['documents'] else [],
                'metadatas': results['metadatas'][0] if results['metadatas'] else [],
                'distances': results['distances'][0] if results['distances'] else [],
                'count': len(results['ids'][0]) if results['ids'] else 0
            }
            
            return formatted_results
            
        except Exception as e:
            logger.error(f"Error searching with embeddings: {e}")
            return {
                'ids': [],
                'documents': [],
                'metadatas': [],
                'distances': [],
                'count': 0,
                'error': str(e)
            }
    
    def get_document(self, doc_id: str) -> Optional[Dict[str, Any]]:
        """
        Get document by ID.
        
        Args:
            doc_id: Document ID
            
        Returns:
            Document data or None
        """
        try:
            results = self.collection.get(
                ids=[doc_id],
                include=["documents", "metadatas"]
            )
            
            if results['ids']:
                return {
                    'id': results['ids'][0],
                    'document': results['documents'][0] if results['documents'] else None,
                    'metadata': results['metadatas'][0] if results['metadatas'] else None
                }
            
            return None
            
        except Exception as e:
            logger.error(f"Error getting document {doc_id}: {e}")
            return None
    
    def update_document(self,
                       doc_id: str,
                       text: Optional[str] = None,
                       metadata: Optional[Dict[str, Any]] = None,
                       embedding: Optional[List[float]] = None) -> bool:
        """
        Update document in vector store.
        
        Args:
            doc_id: Document ID
            text: New text
            metadata: New metadata
            embedding: New embedding
            
        Returns:
            True if successful
        """
        try:
            update_data = {}
            
            if text is not None:
                update_data['documents'] = text
            
            if metadata is not None:
                update_data['metadatas'] = metadata
            
            if embedding is not None:
                update_data['embeddings'] = embedding
            
            if update_data:
                self.collection.update(
                    ids=[doc_id],
                    **update_data
                )
                logger.debug(f"Updated document: {doc_id}")
                return True
            
            return False
            
        except Exception as e:
            logger.error(f"Error updating document {doc_id}: {e}")
            return False
    
    def delete_documents(self, doc_ids: List[str]) -> bool:
        """
        Delete documents from vector store.
        
        Args:
            doc_ids: List of document IDs to delete
            
        Returns:
            True if successful
        """
        try:
            self.collection.delete(ids=doc_ids)
            logger.info(f"Deleted {len(doc_ids)} documents from vector store")
            return True
            
        except Exception as e:
            logger.error(f"Error deleting documents: {e}")
            return False
    
    def get_collection_info(self) -> Dict[str, Any]:
        """
        Get information about the collection.
        
        Returns:
            Collection information
        """
        try:
            count = self.collection.count()
            
            sample = self.collection.get(limit=1)
            
            return {
                'name': self.collection_name,
                'document_count': count,
                'has_documents': count > 0,
                'sample_document': sample['documents'][0] if sample['documents'] else None,
                'persist_directory': str(self.persist_directory)
            }
            
        except Exception as e:
            logger.error(f"Error getting collection info: {e}")
            return {
                'name': self.collection_name,
                'error': str(e)
            }
    
    def reset_collection(self) -> bool:
        """
        Reset/clear the collection.
        
        Returns:
            True if successful
        """
        try:
            self.client.delete_collection(name=self.collection_name)
            self.collection = self._get_or_create_collection()
            logger.info(f"Collection {self.collection_name} has been reset")
            return True
            
        except Exception as e:
            logger.error(f"Error resetting collection: {e}")
            return False
    
    def list_collections(self) -> List[str]:
        """
        List all collections in the database.
        
        Returns:
            List of collection names
        """
        try:
            collections = self.client.list_collections()
            return [col.name for col in collections]
            
        except Exception as e:
            logger.error(f"Error listing collections: {e}")
            return []
    
    def create_new_collection(self, name: str, metadata: Optional[Dict[str, Any]] = None) -> bool:
        """
        Create a new collection.
        
        Args:
            name: Collection name
            metadata: Collection metadata
            
        Returns:
            True if successful
        """
        try:
            self.client.create_collection(
                name=name,
                metadata=metadata or {}
            )
            logger.info(f"Created new collection: {name}")
            return True
            
        except Exception as e:
            logger.error(f"Error creating collection {name}: {e}")
            return False


_chroma_instance = None

def get_chroma_client(
    persist_directory: Optional[Path] = None,
    collection_name: str = "documents"
) -> ChromaClient:
    """
    Get or create ChromaClient instance.
    
    Args:
        persist_directory: Directory to persist ChromaDB data
        collection_name: Name of the collection to use
        
    Returns:
        ChromaClient instance
    """
    global _chroma_instance
    
    if _chroma_instance is None:
        _chroma_instance = ChromaClient(persist_directory, collection_name)
    
    return _chroma_instance


if __name__ == "__main__":
    client = ChromaClient()
    
    test_texts = [
        "The quick brown fox jumps over the lazy dog.",
        "Artificial intelligence is transforming the world.",
        "Python is a popular programming language for data science.",
        "Machine learning algorithms learn from data patterns.",
        "Natural language processing enables computers to understand human language."
    ]
    
    test_metadatas = [
        {"source": "test", "type": "example", "length": len(test_texts[0])},
        {"source": "test", "type": "example", "length": len(test_texts[1])},
        {"source": "test", "type": "example", "length": len(test_texts[2])},
        {"source": "test", "type": "example", "length": len(test_texts[3])},
        {"source": "test", "type": "example", "length": len(test_texts[4])}
    ]
    
    print("Adding test documents...")
    doc_ids = client.add_documents(texts=test_texts, metadatas=test_metadatas)
    print(f"Added {len(doc_ids)} documents")
    
    print("\nSearching for 'artificial intelligence'...")
    results = client.search(query="artificial intelligence", n_results=3)
    print(f"Found {results['count']} results")
    
    for i, (doc, dist) in enumerate(zip(results['documents'], results['distances'])):
        print(f"{i+1}. {doc[:50]}... (distance: {dist:.4f})")
    
    print("\nCollection info:")
    info = client.get_collection_info()
    for key, value in info.items():
        print(f"  {key}: {value}")
'''
        
        file_path.write_text(content)
    
    def complete_llm_client(self):
        """Complete LLM client"""
        file_path = self.project_dir / "src" / "ai_engine" / "llm_client.py"
        
        content = '''"""
Ollama LLM Client for DocuBot
"""

import ollama
from typing import Dict, List, Any, Optional, Generator
import logging
from datetime import datetime

logger = logging.getLogger(__name__)


class LLMClient:
    """Ollama LLM client with streaming support"""
    
    def __init__(self, model: str = "llama2:7b", host: str = "http://localhost:11434"):
        """
        Initialize LLM client.
        
        Args:
            model: Default model to use
            host: Ollama server host
        """
        self.model = model
        self.host = host
        
        ollama._client.Client(host=host)
        
        logger.info(f"LLM client initialized with model: {model}")
    
    def generate(self,
                prompt: str,
                model: Optional[str] = None,
                stream: bool = False,
                options: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """
        Generate response from LLM.
        
        Args:
            prompt: Input prompt
            model: Model to use
            stream: Whether to stream the response
            options: Generation options
            
        Returns:
            Response dictionary
        """
        model = model or self.model
        
        default_options = {
            'temperature': 0.1,
            'top_p': 0.9,
            'top_k': 40,
            'num_predict': 1024,
            'stop': []
        }
        
        if options:
            default_options.update(options)
        
        try:
            if stream:
                response_generator = self._generate_streaming(
                    prompt=prompt,
                    model=model,
                    options=default_options
                )
                return {
                    'success': True,
                    'model': model,
                    'stream': True,
                    'generator': response_generator,
                    'timestamp': datetime.now().isoformat()
                }
            else:
                response = ollama.generate(
                    model=model,
                    prompt=prompt,
                    options=default_options
                )
                
                return {
                    'success': True,
                    'model': model,
                    'response': response['response'],
                    'total_duration': response.get('total_duration', 0),
                    'load_duration': response.get('load_duration', 0),
                    'prompt_eval_count': response.get('prompt_eval_count', 0),
                    'prompt_eval_duration': response.get('prompt_eval_duration', 0),
                    'eval_count': response.get('eval_count', 0),
                    'eval_duration': response.get('eval_duration', 0),
                    'timestamp': datetime.now().isoformat()
                }
                
        except Exception as e:
            logger.error(f"LLM generation failed: {e}")
            return {
                'success': False,
                'error': str(e),
                'model': model,
                'timestamp': datetime.now().isoformat()
            }
    
    def _generate_streaming(self,
                          prompt: str,
                          model: str,
                          options: Dict[str, Any]) -> Generator[str, None, None]:
        """
        Generate streaming response.
        
        Args:
            prompt: Input prompt
            model: Model to use
            options: Generation options
            
        Yields:
            Response chunks
        """
        try:
            stream = ollama.generate(
                model=model,
                prompt=prompt,
                options=options,
                stream=True
            )
            
            for chunk in stream:
                if 'response' in chunk:
                    yield chunk['response']
                    
        except Exception as e:
            logger.error(f"Streaming generation failed: {e}")
            yield f"Error: {str(e)}"
    
    def chat(self,
            messages: List[Dict[str, str]],
            model: Optional[str] = None,
            stream: bool = False,
            options: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """
        Chat completion with conversation history.
        
        Args:
            messages: List of message dictionaries
            model: Model to use
            stream: Whether to stream the response
            options: Generation options
            
        Returns:
            Response dictionary
        """
        model = model or self.model
        
        default_options = {
            'temperature': 0.1,
            'top_p': 0.9,
            'num_predict': 1024
        }
        
        if options:
            default_options.update(options)
        
        try:
            if stream:
                response_generator = self._chat_streaming(
                    messages=messages,
                    model=model,
                    options=default_options
                )
                return {
                    'success': True,
                    'model': model,
                    'stream': True,
                    'generator': response_generator,
                    'timestamp': datetime.now().isoformat()
                }
            else:
                response = ollama.chat(
                    model=model,
                    messages=messages,
                    options=default_options
                )
                
                return {
                    'success': True,
                    'model': model,
                    'response': response['message']['content'],
                    'total_duration': response.get('total_duration', 0),
                    'load_duration': response.get('load_duration', 0),
                    'prompt_eval_count': response.get('prompt_eval_count', 0),
                    'prompt_eval_duration': response.get('prompt_eval_duration', 0),
                    'eval_count': response.get('eval_count', 0),
                    'eval_duration': response.get('eval_duration', 0),
                    'timestamp': datetime.now().isoformat()
                }
                
        except Exception as e:
            logger.error(f"Chat completion failed: {e}")
            return {
                'success': False,
                'error': str(e),
                'model': model,
                'timestamp': datetime.now().isoformat()
            }
    
    def _chat_streaming(self,
                       messages: List[Dict[str, str]],
                       model: str,
                       options: Dict[str, Any]) -> Generator[str, None, None]:
        """
        Streaming chat completion.
        
        Args:
            messages: List of message dictionaries
            model: Model to use
            options: Generation options
            
        Yields:
            Response chunks
        """
        try:
            stream = ollama.chat(
                model=model,
                messages=messages,
                options=options,
                stream=True
            )
            
            for chunk in stream:
                if 'message' in chunk and 'content' in chunk['message']:
                    yield chunk['message']['content']
                    
        except Exception as e:
            logger.error(f"Streaming chat failed: {e}")
            yield f"Error: {str(e)}"
    
    def list_models(self) -> List[Dict[str, Any]]:
        """
        List available models.
        
        Returns:
            List of model information
        """
        try:
            models = ollama.list()
            
            formatted_models = []
            for model in models.get('models', []):
                formatted_models.append({
                    'name': model.get('name', ''),
                    'modified_at': model.get('modified_at', ''),
                    'size': model.get('size', 0),
                    'digest': model.get('digest', ''),
                    'details': model.get('details', {})
                })
            
            return formatted_models
            
        except Exception as e:
            logger.error(f"Failed to list models: {e}")
            return []
    
    def pull_model(self, model_name: str) -> Dict[str, Any]:
        """
        Pull/download a model.
        
        Args:
            model_name: Name of model to pull
            
        Returns:
            Pull status
        """
        try:
            response = ollama.pull(model_name)
            
            return {
                'success': True,
                'model': model_name,
                'status': response.get('status', ''),
                'digest': response.get('digest', ''),
                'total': response.get('total', 0),
                'completed': response.get('completed', 0)
            }
            
        except Exception as e:
            logger.error(f"Failed to pull model {model_name}: {e}")
            return {
                'success': False,
                'model': model_name,
                'error': str(e)
            }
    
    def delete_model(self, model_name: str) -> bool:
        """
        Delete a model.
        
        Args:
            model_name: Name of model to delete
            
        Returns:
            True if successful
        """
        try:
            ollama.delete(model_name)
            logger.info(f"Deleted model: {model_name}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to delete model {model_name}: {e}")
            return False
    
    def get_model_info(self, model_name: Optional[str] = None) -> Dict[str, Any]:
        """
        Get information about a model.
        
        Args:
            model_name: Model name
            
        Returns:
            Model information
        """
        model_name = model_name or self.model
        
        try:
            info = ollama.show(model_name)
            
            return {
                'success': True,
                'model': model_name,
                'license': info.get('license', ''),
                'modelfile': info.get('modelfile', ''),
                'parameters': info.get('parameters', ''),
                'template': info.get('template', ''),
                'details': info.get('details', {})
            }
            
        except Exception as e:
            logger.error(f"Failed to get model info for {model_name}: {e}")
            return {
                'success': False,
                'model': model_name,
                'error': str(e)
            }
    
    def health_check(self) -> Dict[str, Any]:
        """
        Check if Ollama is running and accessible.
        
        Returns:
            Health check results
        """
        try:
            models = self.list_models()
            
            return {
                'success': True,
                'status': 'healthy',
                'host': self.host,
                'available_models': len(models),
                'default_model': self.model,
                'timestamp': datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Ollama health check failed: {e}")
            return {
                'success': False,
                'status': 'unhealthy',
                'host': self.host,
                'error': str(e),
                'timestamp': datetime.now().isoformat()
            }


_llm_instance = None

def get_llm_client(model: str = "llama2:7b", host: str = "http://localhost:11434") -> LLMClient:
    """
    Get or create LLMClient instance.
    
    Args:
        model: Default model to use
        host: Ollama server host
        
    Returns:
        LLMClient instance
    """
    global _llm_instance
    
    if _llm_instance is None:
        _llm_instance = LLMClient(model, host)
    
    return _llm_instance


if __name__ == "__main__":
    client = LLMClient()
    
    print("Performing health check...")
    health = client.health_check()
    print(f"Health: {health['status']}")
    print(f"Available models: {health.get('available_models', 0)}")
    
    if health['success']:
        print("\nTesting generation...")
        response = client.generate(
            prompt="Explain artificial intelligence in one sentence.",
            stream=False
        )
        
        if response['success']:
            print(f"Response: {response['response']}")
            print(f"Duration: {response.get('total_duration', 0) / 1e9:.2f}s")
        
        print("\nAvailable models:")
        models = client.list_models()
        for model in models[:3]:
            print(f"  - {model['name']} ({model['size'] / 1e9:.1f}GB)")
'''
        
        file_path.write_text(content)
    
    def complete_embedding_service(self):
        """Complete embedding service"""
        file_path = self.project_dir / "src" / "ai_engine" / "embedding_service.py"
        
        content = '''"""
Embedding Service for DocuBot
"""

from sentence_transformers import SentenceTransformer
from typing import List, Dict, Any, Optional
import logging
import numpy as np
from pathlib import Path
import pickle
import hashlib

logger = logging.getLogger(__name__)


class EmbeddingService:
    """Service for generating text embeddings"""
    
    def __init__(self, model_name: str = "all-MiniLM-L6-v2", cache_dir: Optional[Path] = None):
        """
        Initialize embedding service.
        
        Args:
            model_name: Name of the embedding model
            cache_dir: Directory for embedding cache
        """
        self.model_name = model_name
        
        logger.info(f"Loading embedding model: {model_name}")
        self.model = SentenceTransformer(model_name)
        logger.info(f"Model loaded: {model_name} (dimensions: {self.model.get_sentence_embedding_dimension()})")
        
        if cache_dir is None:
            from ..core.constants import DATA_DIR
            cache_dir = DATA_DIR / "cache" / "embeddings"
        
        self.cache_dir = cache_dir
        self.cache_dir.mkdir(parents=True, exist_ok=True)
        
        self.cache_index = self._load_cache_index()
        
        logger.info(f"Embedding service initialized with cache: {self.cache_dir}")
    
    def generate_embeddings(self, texts: List[str], batch_size: int = 32) -> List[List[float]]:
        """
        Generate embeddings for a list of texts.
        
        Args:
            texts: List of text strings
            batch_size: Batch size for processing
            
        Returns:
            List of embedding vectors
        """
        if not texts:
            return []
        
        cached_embeddings, uncached_texts = self._get_cached_embeddings(texts)
        
        if not uncached_texts:
            logger.debug(f"All {len(texts)} embeddings retrieved from cache")
            return cached_embeddings
        
        logger.debug(f"Generating embeddings for {len(uncached_texts)} uncached texts")
        
        try:
            new_embeddings = self.model.encode(
                uncached_texts,
                batch_size=batch_size,
                show_progress_bar=False,
                normalize_embeddings=True,
                convert_to_numpy=True
            )
            
            new_embeddings_list = new_embeddings.tolist()
            
            self._cache_embeddings(uncached_texts, new_embeddings_list)
            
            all_embeddings = []
            cache_idx = 0
            new_idx = 0
            
            for text in texts:
                if text in self.cache_index:
                    all_embeddings.append(cached_embeddings[cache_idx])
                    cache_idx += 1
                else:
                    all_embeddings.append(new_embeddings_list[new_idx])
                    new_idx += 1
            
            logger.info(f"Generated {len(uncached_texts)} new embeddings, {len(cached_embeddings)} from cache")
            return all_embeddings
            
        except Exception as e:
            logger.error(f"Error generating embeddings: {e}")
            raise
    
    def generate_single_embedding(self, text: str) -> List[float]:
        """
        Generate embedding for a single text.
        
        Args:
            text: Input text
            
        Returns:
            Embedding vector
        """
        embeddings = self.generate_embeddings([text])
        return embeddings[0] if embeddings else []
    
    def _get_cache_key(self, text: str) -> str:
        """
        Generate cache key for text.
        
        Args:
            text: Input text
            
        Returns:
            Cache key
        """
        hash_input = f"{self.model_name}:{text}"
        return hashlib.md5(hash_input.encode('utf-8')).hexdigest()
    
    def _get_cached_embeddings(self, texts: List[str]) -> tuple:
        """
        Get cached embeddings for texts.
        
        Args:
            texts: List of texts
            
        Returns:
            Tuple of (cached_embeddings, uncached_texts)
        """
        cached_embeddings = []
        uncached_texts = []
        
        for text in texts:
            cache_key = self._get_cache_key(text)
            cache_file = self.cache_dir / f"{cache_key}.pkl"
            
            if cache_file.exists() and cache_key in self.cache_index:
                try:
                    with open(cache_file, 'rb') as f:
                        embedding = pickle.load(f)
                    cached_embeddings.append(embedding)
                except Exception as e:
                    logger.warning(f"Error loading cached embedding: {e}")
                    uncached_texts.append(text)
            else:
                uncached_texts.append(text)
        
        return cached_embeddings, uncached_texts
    
    def _cache_embeddings(self, texts: List[str], embeddings: List[List[float]]):
        """
        Cache embeddings.
        
        Args:
            texts: List of texts
            embeddings: List of embedding vectors
        """
        for text, embedding in zip(texts, embeddings):
            cache_key = self._get_cache_key(text)
            cache_file = self.cache_dir / f"{cache_key}.pkl"
            
            try:
                with open(cache_file, 'wb') as f:
                    pickle.dump(embedding, f)
                
                self.cache_index[cache_key] = {
                    'text_hash': cache_key,
                    'model': self.model_name,
                    'timestamp': __import__('datetime').datetime.now().isoformat()
                }
                
            except Exception as e:
                logger.warning(f"Error caching embedding: {e}")
        
        self._save_cache_index()
    
    def _load_cache_index(self) -> Dict[str, Any]:
        """
        Load cache index.
        
        Returns:
            Cache index dictionary
        """
        index_file = self.cache_dir / "index.pkl"
        
        if index_file.exists():
            try:
                with open(index_file, 'rb') as f:
                    return pickle.load(f)
            except Exception as e:
                logger.warning(f"Error loading cache index: {e}")
        
        return {}
    
    def _save_cache_index(self):
        """Save cache index"""
        index_file = self.cache_dir / "index.pkl"
        
        try:
            with open(index_file, 'wb') as f:
                pickle.dump(self.cache_index, f)
        except Exception as e:
            logger.warning(f"Error saving cache index: {e}")
    
    def clear_cache(self) -> bool:
        """
        Clear embedding cache.
        
        Returns:
            True if successful
        """
        try:
            for cache_file in self.cache_dir.glob("*.pkl"):
                cache_file.unlink()
            
            self.cache_index = {}
            self._save_cache_index()
            
            logger.info(f"Cleared embedding cache: {self.cache_dir}")
            return True
            
        except Exception as e:
            logger.error(f"Error clearing cache: {e}")
            return False
    
    def get_cache_stats(self) -> Dict[str, Any]:
        """
        Get cache statistics.
        
        Returns:
            Cache statistics
        """
        cache_files = list(self.cache_dir.glob("*.pkl"))
        embedding_files = [f for f in cache_files if f.name != "index.pkl"]
        
        total_size = sum(f.stat().st_size for f in cache_files)
        
        return {
            'cache_dir': str(self.cache_dir),
            'total_entries': len(embedding_files),
            'total_size_bytes': total_size,
            'total_size_mb': total_size / (1024 * 1024),
            'model': self.model_name,
            'embedding_dimensions': self.model.get_sentence_embedding_dimension()
        }
    
    def get_model_info(self) -> Dict[str, Any]:
        """
        Get information about the embedding model.
        
        Returns:
            Model information
        """
        return {
            'model_name': self.model_name,
            'embedding_dimensions': self.model.get_sentence_embedding_dimension(),
            'max_seq_length': self.model.max_seq_length,
            'device': str(self.model.device),
            'cache_enabled': True,
            'cache_entries': len(self.cache_index)
        }
    
    def similarity(self, embedding1: List[float], embedding2: List[float]) -> float:
        """
        Calculate cosine similarity between two embeddings.
        
        Args:
            embedding1: First embedding vector
            embedding2: Second embedding vector
            
        Returns:
            Cosine similarity score
        """
        if not embedding1 or not embedding2:
            return 0.0
        
        vec1 = np.array(embedding1)
        vec2 = np.array(embedding2)
        
        similarity = np.dot(vec1, vec2) / (np.linalg.norm(vec1) * np.linalg.norm(vec2))
        
        return float(similarity)
    
    def find_similar(self, query_embedding: List[float], 
                    candidate_embeddings: List[List[float]], 
                    top_k: int = 5) -> List[tuple]:
        """
        Find most similar embeddings.
        
        Args:
            query_embedding: Query embedding
            candidate_embeddings: List of candidate embeddings
            top_k: Number of results to return
            
        Returns:
            List of (index, similarity) tuples
        """
        if not query_embedding or not candidate_embeddings:
            return []
        
        query_vec = np.array(query_embedding)
        candidate_matrix = np.array(candidate_embeddings)
        
        similarities = np.dot(candidate_matrix, query_vec) / (
            np.linalg.norm(candidate_matrix, axis=1) * np.linalg.norm(query_vec)
        )
        
        top_indices = np.argsort(similarities)[::-1][:top_k]
        
        return [(int(idx), float(similarities[idx])) for idx in top_indices]


_embedding_instance = None

def get_embedding_service(model_name: str = "all-MiniLM-L6-v2", 
                         cache_dir: Optional[Path] = None) -> EmbeddingService:
    """
    Get or create EmbeddingService instance.
    
    Args:
        model_name: Name of the embedding model
        cache_dir: Directory for embedding cache
        
    Returns:
        EmbeddingService instance
    """
    global _embedding_instance
    
    if _embedding_instance is None:
        _embedding_instance = EmbeddingService(model_name, cache_dir)
    
    return _embedding_instance


if __name__ == "__main__":
    service = EmbeddingService()
    
    test_texts = [
        "The quick brown fox jumps over the lazy dog.",
        "Artificial intelligence is transforming the world.",
        "Python is a popular programming language."
    ]
    
    print("Generating embeddings...")
    embeddings = service.generate_embeddings(test_texts)
    
    print(f"Generated {len(embeddings)} embeddings")
    print(f"Embedding dimensions: {len(embeddings[0])}")
    
    sim1 = service.similarity(embeddings[0], embeddings[1])
    sim2 = service.similarity(embeddings[0], embeddings[2])
    
    print(f"Similarity between text 1 and 2: {sim1:.4f}")
    print(f"Similarity between text 1 and 3: {sim2:.4f}")
    
    similar = service.find_similar(embeddings[0], embeddings, top_k=2)
    
    for idx, sim in similar:
        print(f"  Text {idx + 1}: similarity = {sim:.4f}")
    
    print("\nCache statistics:")
    stats = service.get_cache_stats()
    for key, value in stats.items():
        print(f"  {key}: {value}")
'''
        
        file_path.write_text(content)
    
    def complete_web_components(self):
        """Complete web components"""
        components_dir = self.project_dir / "src" / "ui" / "web" / "components"
        components_dir.mkdir(parents=True, exist_ok=True)
        
        chat_ui = components_dir / "chat_ui.py"
        chat_ui.write_text('''"""
Web Chat UI Components for Streamlit
"""

import streamlit as st
from typing import List, Dict, Any, Optional
import json
from datetime import datetime


class ChatUI:
    """Chat interface components for Streamlit"""
    
    def __init__(self):
        pass
    
    def display_chat_message(self, role: str, content: str, timestamp: Optional[str] = None):
        """
        Display a chat message.
        
        Args:
            role: Message role
            content: Message content
            timestamp: Optional timestamp
        """
        if role == "user":
            with st.chat_message("user"):
                st.markdown(content)
                if timestamp:
                    st.caption(f"User - {timestamp}")
        else:
            with st.chat_message("assistant"):
                st.markdown(content)
                if timestamp:
                    st.caption(f"Assistant - {timestamp}")
    
    def display_chat_history(self, messages: List[Dict[str, Any]]):
        """
        Display chat history.
        
        Args:
            messages: List of message dictionaries
        """
        for message in messages:
            self.display_chat_message(
                role=message.get('role', 'user'),
                content=message.get('content', ''),
                timestamp=message.get('timestamp')
            )
    
    def chat_input(self, placeholder: str = "Type your message...") -> Optional[str]:
        """
        Display chat input.
        
        Args:
            placeholder: Input placeholder text
            
        Returns:
            User input or None
        """
        return st.chat_input(placeholder)
    
    def display_sources(self, sources: List[Dict[str, Any]]):
        """
        Display source citations.
        
        Args:
            sources: List of source dictionaries
        """
        if sources:
            with st.expander("Sources", expanded=False):
                for i, source in enumerate(sources, 1):
                    st.markdown(f"**Source {i}**")
                    if 'title' in source:
                        st.markdown(f"**Title:** {source['title']}")
                    if 'content' in source:
                        st.markdown(f"**Excerpt:** {source['content'][:200]}...")
                    if 'similarity' in source:
                        st.markdown(f"**Relevance:** {source['similarity']:.2%}")
                    st.divider()
    
    def display_processing_status(self, status: str, message: str = ""):
        """
        Display processing status.
        
        Args:
            status: Status
            message: Status message
        """
        if status == "processing":
            with st.spinner(message or "Processing..."):
                pass
        elif status == "success":
            st.success(message or "Success")
        elif status == "error":
            st.error(message or "Error occurred")
    
    def create_sidebar(self):
        """Create application sidebar"""
        with st.sidebar:
            st.title("Settings")
            
            model = st.selectbox(
                "Model",
                ["llama2:7b", "mistral:7b", "neural-chat:7b"],
                index=0
            )
            
            temperature = st.slider(
                "Temperature",
                min_value=0.0,
                max_value=1.0,
                value=0.1,
                step=0.05,
                help="Higher values make output more random"
            )
            
            st.divider()
            st.subheader("Conversations")
            
            if st.button("New Conversation", use_container_width=True):
                st.session_state.messages = []
                st.rerun()
            
            if st.button("Clear History", use_container_width=True):
                st.session_state.conversation_history = []
                st.rerun()
            
            return {
                'model': model,
                'temperature': temperature
            }
    
    def display_document_upload(self):
        """
        Display document upload interface.
        
        Returns:
            Uploaded files or None
        """
        st.subheader("Upload Documents")
        
        uploaded_files = st.file_uploader(
            "Choose files",
            type=['pdf', 'txt', 'docx', 'md', 'html'],
            accept_multiple_files=True,
            help="Upload documents to add to knowledge base"
        )
        
        if uploaded_files:
            st.success(f"Selected {len(uploaded_files)} file(s)")
            
            for file in uploaded_files:
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.text(f"{file.name} ({file.size / 1024:.1f} KB)")
                with col2:
                    st.button("Process", key=f"process_{file.name}")
        
        return uploaded_files
    
    def display_document_list(self, documents: List[Dict[str, Any]]):
        """
        Display list of documents.
        
        Args:
            documents: List of document dictionaries
        """
        st.subheader("Documents")
        
        if not documents:
            st.info("No documents uploaded yet")
            return
        
        for doc in documents:
            with st.expander(f"{doc.get('file_name', 'Unknown')}", expanded=False):
                col1, col2 = st.columns([3, 1])
                
                with col1:
                    st.markdown(f"**Type:** {doc.get('file_type', 'Unknown')}")
                    st.markdown(f"**Size:** {doc.get('file_size', 0) / 1024:.1f} KB")
                    st.markdown(f"**Status:** {doc.get('processing_status', 'Unknown')}")
                    st.markdown(f"**Chunks:** {doc.get('chunk_count', 0)}")
                
                with col2:
                    if st.button("Delete", key=f"delete_{doc.get('id')}"):
                        st.warning(f"Delete {doc.get('file_name')}?")
    
    def display_statistics(self, stats: Dict[str, Any]):
        """
        Display application statistics.
        
        Args:
            stats: Statistics dictionary
        """
        st.subheader("Statistics")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("Documents", stats.get('total_documents', 0))
        
        with col2:
            st.metric("Chunks", stats.get('total_chunks', 0))
        
        with col3:
            st.metric("Conversations", stats.get('total_conversations', 0))
        
        if 'database_size_mb' in stats:
            st.metric("Database Size", f"{stats['database_size_mb']:.1f} MB")
    
    def display_error(self, error: Exception):
        """
        Display error message.
        
        Args:
            error: Exception object
        """
        st.error(f"Error: {str(error)}")
        
        with st.expander("Error Details", expanded=False):
            st.code(str(error))
    
    def display_info_message(self, message: str):
        """
        Display informational message.
        
        Args:
            message: Message text
        """
        st.info(message)
    
    def display_warning_message(self, message: str):
        """
        Display warning message.
        
        Args:
            message: Message text
        """
        st.warning(message)
    
    def display_success_message(self, message: str):
        """
        Display success message.
        
        Args:
            message: Message text
        """
        st.success(message)


_chat_ui = None

def get_chat_ui() -> ChatUI:
    """
    Get or create ChatUI instance.
    
    Returns:
        ChatUI instance
    """
    global _chat_ui
    
    if _chat_ui is None:
        _chat_ui = ChatUI()
    
    return _chat_ui
''')
        
        document_list = components_dir / "document_list.py"
        document_list.write_text('''"""
Document List Components for Streamlit
"""

import streamlit as st
from typing import List, Dict, Any
import pandas as pd
from datetime import datetime


class DocumentListUI:
    """Document list interface components"""
    
    def __init__(self):
        pass
    
    def display_document_table(self, documents: List[Dict[str, Any]]):
        """
        Display documents in a table.
        
        Args:
            documents: List of document dictionaries
        """
        if not documents:
            st.info("No documents available")
            return
        
        df_data = []
        for doc in documents:
            df_data.append({
                'Name': doc.get('file_name', 'Unknown'),
                'Type': doc.get('file_type', 'Unknown'),
                'Size (KB)': doc.get('file_size', 0) / 1024,
                'Status': doc.get('processing_status', 'Unknown'),
                'Chunks': doc.get('chunk_count', 0),
                'Uploaded': doc.get('upload_date', '')[:10] if doc.get('upload_date') else ''
            })
        
        df = pd.DataFrame(df_data)
        
        st.dataframe(
            df,
            use_container_width=True,
            hide_index=True,
            column_config={
                'Size (KB)': st.column_config.NumberColumn(format="%.1f KB"),
                'Chunks': st.column_config.NumberColumn(format="%d")
            }
        )
    
    def display_document_filters(self):
        """
        Display document filters.
        
        Returns:
            Filter criteria
        """
        col1, col2, col3 = st.columns(3)
        
        with col1:
            file_type = st.multiselect(
                "File Type",
                ['PDF', 'DOCX', 'TXT', 'MD', 'HTML'],
                default=[]
            )
        
        with col2:
            status = st.multiselect(
                "Status",
                ['pending', 'processing', 'completed', 'failed'],
                default=['completed']
            )
        
        with col3:
            sort_by = st.selectbox(
                "Sort By",
                ['upload_date', 'file_name', 'file_size', 'chunk_count'],
                index=0
            )
        
        return {
            'file_type': file_type,
            'status': status,
            'sort_by': sort_by
        }
    
    def display_document_actions(self, document_id: str):
        """
        Display document action buttons.
        
        Args:
            document_id: Document ID
            
        Returns:
            Action performed
        """
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("View", key=f"view_{document_id}", use_container_width=True):
                return "view"
        
        with col2:
            if st.button("Delete", key=f"delete_{document_id}", use_container_width=True):
                return "delete"
        
        with col3:
            if st.button("Reprocess", key=f"reprocess_{document_id}", use_container_width=True):
                return "reprocess"
        
        return None
    
    def display_document_preview(self, document: Dict[str, Any]):
        """
        Display document preview.
        
        Args:
            document: Document dictionary
        """
        with st.expander("Document Details", expanded=True):
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown(f"**File Name:** {document.get('file_name', 'Unknown')}")
                st.markdown(f"**File Type:** {document.get('file_type', 'Unknown')}")
                st.markdown(f"**File Size:** {document.get('file_size', 0) / 1024:.1f} KB")
            
            with col2:
                st.markdown(f"**Status:** {document.get('processing_status', 'Unknown')}")
                st.markdown(f"**Chunks:** {document.get('chunk_count', 0)}")
                st.markdown(f"**Words:** {document.get('word_count', 0)}")
            
            if document.get('metadata'):
                st.divider()
                st.subheader("Metadata")
                st.json(document.get('metadata', {}))
            
            if document.get('tags'):
                st.divider()
                st.subheader("Tags")
                tags = document.get('tags', [])
                tag_chips = " ".join([f"`{tag}`" for tag in tags])
                st.markdown(tag_chips)
            
            if document.get('summary'):
                st.divider()
                st.subheader("Summary")
                st.markdown(document.get('summary', ''))
    
    def display_batch_operations(self):
        """
        Display batch operation controls.
        
        Returns:
            Batch operation to perform
        """
        st.subheader("Batch Operations")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("Select All", use_container_width=True):
                return "select_all"
        
        with col2:
            if st.button("Process Selected", use_container_width=True):
                return "process_selected"
        
        with col3:
            if st.button("Delete Selected", use_container_width=True):
                return "delete_selected"
        
        return None
    
    def display_upload_progress(self, current: int, total: int, filename: str):
        """
        Display upload progress.
        
        Args:
            current: Current file number
            total: Total files
            filename: Current filename
        """
        progress = current / total
        st.progress(progress, text=f"Uploading {filename} ({current}/{total})")
    
    def display_processing_progress(self, document_name: str, step: str, progress: float):
        """
        Display document processing progress.
        
        Args:
            document_name: Document name
            step: Current processing step
            progress: Progress
        """
        st.progress(progress, text=f"Processing {document_name}: {step}")
    
    def display_document_stats(self, stats: Dict[str, Any]):
        """
        Display document statistics.
        
        Args:
            stats: Statistics dictionary
        """
        st.subheader("Document Statistics")
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Total", stats.get('total', 0))
        
        with col2:
            st.metric("Processed", stats.get('processed', 0))
        
        with col3:
            st.metric("Pending", stats.get('pending', 0))
        
        with col4:
            st.metric("Failed", stats.get('failed', 0))
        
        if 'by_type' in stats:
            st.divider()
            st.subheader("By File Type")
            
            type_data = stats['by_type']
            df = pd.DataFrame(list(type_data.items()), columns=['Type', 'Count'])
            st.bar_chart(df.set_index('Type'))
    
    def display_search_bar(self):
        """
        Display document search bar.
        
        Returns:
            Search query
        """
        return st.text_input(
            "Search Documents",
            placeholder="Search by filename, content, or tags...",
            key="document_search"
        )
    
    def display_empty_state(self, message: str = "No documents found"):
        """
        Display empty state.
        
        Args:
            message: Message to display
        """
        st.markdown(f"""
        <div style='text-align: center; padding: 40px;'>
            <h3>{message}</h3>
            <p>Upload documents to get started</p>
        </div>
        """, unsafe_allow_html=True)


_doc_list_ui = None

def get_document_list_ui() -> DocumentListUI:
    """
    Get or create DocumentListUI instance.
    
    Returns:
        DocumentListUI instance
    """
    global _doc_list_ui
    
    if _doc_list_ui is None:
        _doc_list_ui = DocumentListUI()
    
    return _doc_list_ui
''')
    
    def complete_cli_formatters(self):
        """Complete CLI output formatters"""
        file_path = self.project_dir / "src" / "ui" / "cli" / "output_formatters.py"
        
        content = '''"""
CLI Output Formatters for DocuBot
"""

from typing import List, Dict, Any, Optional
from datetime import datetime
import json
from textwrap import wrap
import sys


class CLIFormatter:
    """Command Line Interface output formatter"""
    
    def __init__(self, width: int = 80):
        """
        Initialize CLI formatter.
        
        Args:
            width: Terminal width for formatting
        """
        self.width = width
    
    def format_header(self, title: str, char: str = "=") -> str:
        """
        Format section header.
        
        Args:
            title: Header title
            char: Character to use for line
            
        Returns:
            Formatted header
        """
        line = char * self.width
        centered = title.center(self.width)
        return f"{line}\n{centered}\n{line}"
    
    def format_subheader(self, title: str, char: str = "-") -> str:
        """
        Format subheader.
        
        Args:
            title: Subheader title
            char: Character to use for line
            
        Returns:
            Formatted subheader
        """
        return f"{char * 40} {title} {char * (self.width - 42 - len(title))}"
    
    def format_table(self, headers: List[str], rows: List[List[Any]]) -> str:
        """
        Format data as table.
        
        Args:
            headers: Column headers
            rows: Table rows
            
        Returns:
            Formatted table
        """
        if not headers or not rows:
            return ""
        
        col_widths = []
        for i, header in enumerate(headers):
            max_width = len(str(header))
            for row in rows:
                if i < len(row):
                    max_width = max(max_width, len(str(row[i])))
            col_widths.append(min(max_width + 2, 30))
        
        lines = []
        
        header_line = "| " + " | ".join(
            str(header).ljust(width) for header, width in zip(headers, col_widths)
        ) + " |"
        lines.append(header_line)
        
        separator = "+" + "+".join("-" * (width + 2) for width in col_widths) + "+"
        lines.append(separator)
        
        for row in rows:
            row_line = "| " + " | ".join(
                str(cell)[:width].ljust(width) for cell, width in zip(row, col_widths)
            ) + " |"
            lines.append(row_line)
        
        return "\n".join(lines)
    
    def format_list(self, items: List[Any], bullet: str = "*") -> str:
        """
        Format list.
        
        Args:
            items: List items
            bullet: Bullet character
            
        Returns:
            Formatted list
        """
        return "\n".join(f"  {bullet} {item}" for item in items)
    
    def format_key_value(self, key: str, value: Any, indent: int = 2) -> str:
        """
        Format key-value pair.
        
        Args:
            key: Key name
            value: Value
            indent: Indentation spaces
            
        Returns:
            Formatted key-value
        """
        indent_str = " " * indent
        value_str = str(value)
        
        if len(value_str) > self.width - indent - len(key) - 4:
            wrapped = wrap(value_str, width=self.width - indent - 4)
            value_str = "\n" + "\n".join(f"{indent_str}    {line}" for line in wrapped)
        
        return f"{indent_str}{key}: {value_str}"
    
    def format_dict(self, data: Dict[str, Any], indent: int = 2) -> str:
        """
        Format dictionary.
        
        Args:
            data: Dictionary to format
            indent: Indentation spaces
            
        Returns:
            Formatted dictionary
        """
        lines = []
        for key, value in data.items():
            if isinstance(value, dict):
                lines.append(f"{' ' * indent}{key}:")
                lines.append(self.format_dict(value, indent + 2))
            elif isinstance(value, list):
                lines.append(f"{' ' * indent}{key}:")
                for item in value:
                    lines.append(f"{' ' * (indent + 2)}* {item}")
            else:
                lines.append(self.format_key_value(key, value, indent))
        
        return "\n".join(lines)
    
    def format_json(self, data: Any, indent: int = 2) -> str:
        """
        Format as JSON.
        
        Args:
            data: Data to format
            indent: JSON indentation
            
        Returns:
            Formatted JSON
        """
        return json.dumps(data, indent=indent, ensure_ascii=False)
    
    def format_progress_bar(self, current: int, total: int, width: int = 30) -> str:
        """
        Format progress bar.
        
        Args:
            current: Current progress
            total: Total steps
            width: Progress bar width
            
        Returns:
            Formatted progress bar
        """
        if total == 0:
            return "[--------------------] 0/0 (100%)"
        
        percentage = current / total
        filled = int(width * percentage)
        bar = "#" * filled + "-" * (width - filled)
        
        return f"[{bar}] {current}/{total} ({percentage:.1%})"
    
    def format_duration(self, seconds: float) -> str:
        """
        Format duration.
        
        Args:
            seconds: Duration in seconds
            
        Returns:
            Formatted duration
        """
        if seconds < 1:
            return f"{seconds * 1000:.0f}ms"
        elif seconds < 60:
            return f"{seconds:.1f}s"
        elif seconds < 3600:
            minutes = seconds / 60
            return f"{minutes:.1f}m"
        else:
            hours = seconds / 3600
            return f"{hours:.1f}h"
    
    def format_file_size(self, bytes_size: int) -> str:
        """
        Format file size.
        
        Args:
            bytes_size: Size in bytes
            
        Returns:
            Formatted file size
        """
        for unit in ['B', 'KB', 'MB', 'GB']:
            if bytes_size < 1024.0:
                return f"{bytes_size:.1f} {unit}"
            bytes_size /= 1024.0
        return f"{bytes_size:.1f} TB"
    
    def format_timestamp(self, timestamp: Optional[str] = None) -> str:
        """
        Format timestamp.
        
        Args:
            timestamp: ISO format timestamp
            
        Returns:
            Formatted timestamp
        """
        if timestamp:
            dt = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
        else:
            dt = datetime.now()
        
        return dt.strftime("%Y-%m-%d %H:%M:%S")
    
    def format_document_info(self, document: Dict[str, Any]) -> str:
        """
        Format document information.
        
        Args:
            document: Document dictionary
            
        Returns:
            Formatted document info
        """
        lines = []
        
        lines.append(self.format_header("Document Information"))
        lines.append(f"Name: {document.get('file_name', 'Unknown')}")
        lines.append(f"Type: {document.get('file_type', 'Unknown')}")
        lines.append(f"Size: {self.format_file_size(document.get('file_size', 0))}")
        lines.append(f"Status: {document.get('processing_status', 'Unknown')}")
        lines.append(f"Chunks: {document.get('chunk_count', 0)}")
        lines.append(f"Uploaded: {document.get('upload_date', '')[:10]}")
        
        if document.get('processing_error'):
            lines.append(f"Error: {document.get('processing_error')}")
        
        return "\n".join(lines)
    
    def format_search_results(self, results: List[Dict[str, Any]]) -> str:
        """
        Format search results.
        
        Args:
            results: Search results
            
        Returns:
            Formatted results
        """
        if not results:
            return "No results found"
        
        lines = []
        lines.append(self.format_header(f"Search Results ({len(results)} found)"))
        
        for i, result in enumerate(results, 1):
            lines.append(f"{i}. {result.get('file_name', 'Unknown')}")
            lines.append(f"   Type: {result.get('file_type', 'Unknown')}")
            lines.append(f"   Score: {result.get('similarity', 0):.3f}")
            
            if 'excerpt' in result:
                excerpt = result['excerpt']
                if len(excerpt) > 100:
                    excerpt = excerpt[:100] + "..."
                lines.append(f"   Excerpt: {excerpt}")
            
            lines.append("")
        
        return "\n".join(lines)
    
    def format_statistics(self, stats: Dict[str, Any]) -> str:
        """
        Format statistics.
        
        Args:
            stats: Statistics dictionary
            
        Returns:
            Formatted statistics
        """
        lines = []
        lines.append(self.format_header("Statistics"))
        
        lines.append(f"Documents: {stats.get('total_documents', 0)}")
        lines.append(f"Chunks: {stats.get('total_chunks', 0)}")
        lines.append(f"Conversations: {stats.get('total_conversations', 0)}")
        
        if 'database_size_bytes' in stats:
            lines.append(f"Database Size: {self.format_file_size(stats['database_size_bytes'])}")
        
        if 'documents_by_status' in stats:
            lines.append("")
            lines.append(self.format_subheader("Document Status"))
            for status, count in stats['documents_by_status'].items():
                lines.append(f"  {status}: {count}")
        
        return "\n".join(lines)
    
    def print_colored(self, text: str, color: str = "default") -> None:
        """
        Print colored text.
        
        Args:
            text: Text to print
            color: Color name
        """
        colors = {
            'red': '\033[91m',
            'green': '\033[92m',
            'yellow': '\033[93m',
            'blue': '\033[94m',
            'magenta': '\033[95m',
            'cyan': '\033[96m',
            'white': '\033[97m',
            'default': '\033[0m'
        }
        
        end_color = '\033[0m'
        
        if color in colors and sys.stdout.isatty():
            print(f"{colors[color]}{text}{end_color}")
        else:
            print(text)


_cli_formatter = None

def get_cli_formatter(width: int = 80) -> CLIFormatter:
    """
    Get or create CLIFormatter instance.
    
    Args:
        width: Terminal width
        
    Returns:
        CLIFormatter instance
    """
    global _cli_formatter
    
    if _cli_formatter is None:
        _cli_formatter = CLIFormatter(width)
    
    return _cli_formatter


if __name__ == "__main__":
    formatter = CLIFormatter()
    
    print(formatter.format_header("DocuBot CLI"))
    print()
    
    headers = ["ID", "Name", "Status", "Size"]
    rows = [
        ["1", "document.pdf", "processed", "1.2 MB"],
        ["2", "notes.txt", "pending", "45 KB"],
        ["3", "report.docx", "failed", "2.5 MB"]
    ]
    
    print(formatter.format_table(headers, rows))
    print()
    
    print(formatter.format_key_value("Version", "1.0.0"))
    print(formatter.format_key_value("Model", "llama2:7b"))
    print()
    
    stats = {
        'total_documents': 15,
        'total_chunks': 125,
        'total_conversations': 8,
        'database_size_bytes': 1024 * 1024 * 50,
        'documents_by_status': {
            'completed': 12,
            'pending': 2,
            'failed': 1
        }
    }
    
    print(formatter.format_statistics(stats))
'''
        
        file_path.write_text(content)
    
    def complete_app_config(self):
        """Enhance app_config.yaml"""
        config_file = self.project_dir / "data" / "config" / "app_config.yaml"
        
        if config_file.exists():
            content = config_file.read_text()
            
            if 'database:' not in content:
                content += '''
database:
  auto_backup: true
  backup_interval_hours: 24
  max_backup_files: 10
  vacuum_on_startup: true

search:
  default_top_k: 5
  similarity_threshold: 0.7
  enable_hybrid_search: true
  rerank_enabled: false

logging:
  level: "INFO"
  file_logging: true
  max_log_size_mb: 10
  backup_count: 5

features:
  enable_ocr: false
  enable_web_extraction: false
  enable_auto_tagging: true
  enable_auto_summarization: true
'''
            
            config_file.write_text(content)


def main():
    """Main function"""
    import argparse
    
    parser = argparse.ArgumentParser(description="Complete DocuBot code")
    parser.add_argument("--dir", default="DocuBot", help="Project directory")
    
    args = parser.parse_args()
    
    completer = CodeCompleter(args.dir)
    completer.complete_all_code()


if __name__ == "__main__":
    main()